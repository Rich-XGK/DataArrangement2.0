"""
    Author:Jack Xu
    Gmail:jack2919048985@gmail.com
"""
import os
import os.path as path
import re
import time
from copy import deepcopy
from typing import List, Dict, Union

import docx as dx
import pandas as pd
import win32com.client as win32
import xlsxwriter


def doc_to_docx(file: str) -> str:
    """
    将.doc文档转换为.docx文档
    :param file: 文件名（绝对路径）
    :return: 转换后的文件名（绝对路径）
    """
    word = win32.Dispatch("Word.Application")
    doc = word.Documents.Open(file)
    file_new = file.replace('.doc', '.docx')
    if not path.exists(file_new):
        doc.SaveAs(file_new, 12)
    doc.Close()
    word.Quit()
    time.sleep(1)
    os.remove(file)
    return file_new


def wps_to_docx(file: str) -> str:
    """
    将.wps文档转换为.docx文档
    :param file: 文件名（绝对路径）
    :return: 转换后的文件名（绝对路径）
    """
    word = win32.Dispatch("Word.Application")
    doc = word.Documents.Open(file)
    file_new = file.replace('.wps', '.docx')
    if not path.exists(file_new):
        doc.SaveAs(file_new, 12)
    doc.Close()
    word.Quit()
    time.sleep(1)
    os.remove(file)
    return file_new


def xls_to_xlsx(file: str) -> str:
    """
    将.xls表格转换为.xlsx表格
    :param file: 文件名（绝对路径）
    :return: 转换后的文件名（绝对路径）
    """
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    wb = excel.Workbooks.Open(file)
    file_new = file.replace('.xls', '.xlsx')
    if not path.exists(file_new):
        wb.SaveAs(file_new, FileFormat=51)
    wb.Close()
    excel.Application.Quit()
    os.remove(file)
    return file_new


def docx01_or_docx02(file: str) -> int:
    """
    判断.docx文档是附件一还是附件二
    :param file: 文件名（绝对路径）
    :return: 1 -> 附件1 ; 2 -> 附件二 ; 3 -> 异常
    """
    try:
        if file.endswith('.doc'):
            file = doc_to_docx(file)
        elif file.endswith('.wps'):
            file = wps_to_docx(file)
        docx = dx.Document(file)
        if len(docx.tables) == 0:
            return 3
        elif len(docx.tables) == 1 or len(docx.tables) == 2:
            if docx.tables[0].cell(0, 0).text == "编号" or docx.tables[0].cell(1, 0).text == "建筑年代":
                return 1
            elif docx.tables[0].cell(0, 0).text == "自然村地址" or docx.tables[0].cell(1, 0).text == "房屋总数（栋）":
                return 2
        else:
            return 1
    except Exception:
        return 3


def xlsx01_or_xlsx02(file: str) -> bool:
    """
    判断.xlsx表格是单体抗震调查表还是整体抗震统计表
    :param file: 文件名（绝对路径）
    :return: True -> 单体抗震调查表 ; False -> 整体抗震统计表
    """
    excel_df = pd.read_excel(file)
    if excel_df.shape[1] >= 20:
        return True
    else:
        return False


def get_excel01_dict(location: Dict[str, str], name: str, phone: str, table) -> Dict[str, str]:
    """
    从附件1中获取信息
    :param location: 包含户所在的区名、镇名、村名的字典
    :param name: 该户姓名
    :param phone: 该户联系方式
    :param table: 从该户 word01 文档中解析得到的table对象
    :return: 包含 word01 文档中重要信息的字典
    """
    info_dic = dict()
    info_dic["编号"] = table.cell(0, 2).text

    text = table.cell(28, 5).text
    text = re.findall(r'\d{1,4}', text)
    info_dic["北纬"] = text[0] + '°' + text[1] + '′' if len(text) == 4 else ""
    info_dic["东经"] = text[2] + '°' + text[3] + '′' if len(text) == 4 else ""

    info_dic["区"] = location["region_name"]
    info_dic["乡镇"] = location["town_name"]
    info_dic["村庄"] = location["village_name"]
    info_dic["户主"] = name
    info_dic["户主联系方式"] = phone
    info_dic["建筑年代"] = table.cell(1, 2).text
    info_dic["层数"] = table.cell(1, 8).text
    info_dic["结构类型"] = table.cell(1, 18).text
    info_dic["建筑高度"] = table.cell(2, 2).text
    info_dic["建筑宽度"] = table.cell(2, 8).text
    info_dic["建筑长度"] = table.cell(2, 18).text
    info_dic["常住人口"] = table.cell(3, 18).text

    text = table.cell(4, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{2,4} {0,2}√", text)
        info_dic["平面"] = text[0].replace('√', '') if text else ""
        info_dic["立面"] = text[1].replace('√', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text01_ls, text02_ls = re.findall(r'平面:不?规则', text), re.findall(r'立面:不?规则', text)
        info_dic["平面"] = text01_ls[0].split(':')[-1] if text01_ls else ''
        info_dic["立面"] = text02_ls[0].split(':')[-1] if text02_ls else ''

    info_dic["房屋间数"] = table.cell(4, 18).text

    text = table.cell(5, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{2,4} {0,2}√", text)
        info_dic["场地条件"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'场地条件:[\u4E00-\u9FA5]{2,4}', text)
        info_dic["场地条件"] = text[0].split(':')[-1] if text else ''

    text = table.cell(6, 3).text
    if text.find('√') != -1:
        text = re.findall(r"\d+\.?\d*m {0,2}√", text)
        info_dic["基坑深度"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'基坑深度[(]m[)]:\d[.]?\d?m', text)
        info_dic["基坑深度"] = text[0].split(':')[-1] if text else ''

    text = table.cell(7, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{2,3} {0,2}√", text)
        info_dic["基坑回填材料"] = text[0].replace('√', '').replace(' ', '') if text else "三合土（石灰+黏土+细砂）"
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'基坑回填材料:.*$', text)
        info_dic["基坑回填材料"] = text[0].split(':')[-1] if text else ''

    text = table.cell(8, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,3} {0,2}√", text)
        info_dic["基础材料"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'基础材料:[\u4E00-\u9FA5]{1,3}', text)
        info_dic["基础材料"] = text[0].split(':')[-1] if text else ''

    text = table.cell(9, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,4} {0,2}√", text)
        info_dic["基础砌筑砂浆材料"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'基础砌筑砂浆材料:[\u4E00-\u9FA5]{1,4}', text)
        info_dic["基础砌筑砂浆材料"] = text[0].split(':')[-1] if text else ''

    text = table.cell(12, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,4} {0,2}√", text)
        info_dic["墙体砌块材料"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'砌块材料:[\u4E00-\u9FA5]{2}', text)
        info_dic["墙体砌块材料"] = text[0].split(':')[-1] if text else ''

    text = table.cell(13, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,4} {0,2}√", text)
        info_dic["墙体砂浆材料"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'砂浆材料:[\u4E00-\u9FA5]{1,4}', text)
        info_dic["墙体砂浆材料"] = text[0].split(':')[-1] if text else ''

    text = table.cell(14, 3).text
    if text.find('√') != -1:
        text = re.findall(r"\d{2,4}", text)
        info_dic["外墙厚度"] = text[0] if text else ""
        info_dic["内墙厚度"] = text[1] if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text01_ls, text02_ls = re.findall(r'外墙\d{1,2}', text), re.findall(r'内墙(房间隔断墙)\d{1,2}', text)
        info_dic["外墙厚度"] = text01_ls[0].replace('外墙', '') + 'cm' if text01_ls else ""
        info_dic["内墙厚度"] = text02_ls[0].replace('内墙(房间隔断墙)', '') + 'cm' if text02_ls else ""

    text = table.cell(15, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5] {0,2}√", text)
        info_dic["烟道"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'墙体内竖向孔道[(]烟囱道[)]:[有无]', text)
        info_dic["烟道"] = text[0].split(':')[-1] if text else ''

    text = table.cell(16, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5] {0,2}√", text)
        info_dic["女儿墙"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'女儿墙[(]屋顶周围的矮墙[)]:[有无]', text)
        info_dic["女儿墙"] = text[0].split(':')[-1] if text else ''

    info_dic["上部圈梁"] = "有"  # TODO
    info_dic["上部圈梁闭合"] = "基本"
    info_dic["基础圈梁"] = "有"
    info_dic["基础圈梁闭合"] = "基本"

    text = table.cell(20, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{5} {0,2}√", text)
        info_dic["构造柱"] = "未设置" if text[0].find("未设") != -1 else "设置"
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        info_dic["构造柱"] = "未设置" if text.find("未设") != -1 else "设置"

    text = table.cell(21, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,5}.{5} {0,2}√", text)
        if text:
            info_dic["屋盖类别"] = "坡顶房" if text[0].find("坡顶房") != -1 else "平顶房"
        else:
            info_dic["屋盖类别"] = ''
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'屋盖类别:[\u4E00-\u9FA5]{2,5}[(][坡平]顶房[)]', text)
        if text:
            info_dic["屋盖类别"] = "坡顶房" if text[0].find("坡顶房") != -1 else "平顶房"
        else:
            info_dic["屋盖类别"] = ''

    text = table.cell(22, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{1,3} {0,2}√", text)
        info_dic["楼板类别"] = text[0].replace('√', '').replace(' ', '') if text else ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'楼板类别:[\u4E00-\u9FA5]{1,3}', text)
        info_dic["楼板类别"] = text[0].split(':')[-1] if text else ''

    text = table.cell(23, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{3} {0,2}√", text)
        info_dic["墙体歪闪"] = "无" if text[0].find("无") != -1 else "有"
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'墙体:[\u4E00-\u9FA5]{1,3}', text)
        if text:
            info_dic["墙体歪闪"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["墙体歪闪"] = ""

    text = table.cell(23, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]腐蚀、酥碎 {0,2}√", text)
        if text:
            info_dic["墙体腐蚀"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["墙体腐蚀"] = ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'墙体:[\u4E00-\u9FA5]{1,3}、?[\u4E00-\u9FA5]{0,2}', text)
        if text:
            info_dic["墙体腐蚀"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["墙体腐蚀"] = ""

    text = table.cell(24, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{5,7} {0,2}√", text)
        if text:
            info_dic["墙体裂缝"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["墙体裂缝"] = ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'墙体:[\u4E00-\u9FA5]{5,7}', text)
        if text:
            info_dic["墙体裂缝"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["墙体裂缝"] = ""

    text = table.cell(25, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]{5,7} {0,2}√", text)
        if text:
            info_dic["基础沉降"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["基础沉降"] = ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'基础:[\u4E00-\u9FA5]{6}', text)
        if text:
            info_dic["基础沉降"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["基础沉降"] = ""

    text = table.cell(26, 3).text
    if text.find('√') != -1:
        text = re.findall(r"[\u4E00-\u9FA5]变形、腐朽或开裂 {0,2}√", text)
        if text:
            info_dic["屋盖情况"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["屋盖情况"] = ""
    else:
        text = text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace('（', '(').replace('）', ')').replace('\n', '')
        text = re.findall(r'楼、屋盖构件:[\u4E00-\u9FA5]{3}、[\u4E00-\u9FA5]{5}', text)
        if text:
            info_dic["屋盖情况"] = "无" if text[0].find("无") != -1 else "有"
        else:
            info_dic["屋盖情况"] = ""

    info_dic["填表人"] = table.cell(29, 3).text
    info_dic["填表人联系方式"] = table.cell(29, 9).text

    return info_dic


def get_excel02_dict(table) -> Dict[str, str]:
    """
    从附件2中获取信息
    :param table: 从该户 word02 文档中解析得到的table对
    :return: 包含 word02 文档中重要信息的字典
    """
    table_rows = len(table.rows)
    info_dic = {}
    for i in range(22):
        flag = True if i < table_rows else False
        if i == 0:
            info_dic["自然村地址"] = table.cell(0, 1).text if flag else ""
            info_dic["住户总数（户）"] = table.cell(0, 5).text if flag else ""
            info_dic["人口总数（口）"] = table.cell(0, 8).text if flag else ""
        elif i == 1:
            info_dic["房屋总数（栋）"] = table.cell(1, 1).text if flag else ""
            info_dic["家庭平均人口（口）"] = table.cell(1, 3).text if flag else ""
            info_dic["全村上一年经济收入（万元）"] = table.cell(1, 8).text if flag else ""
        elif i == 2:
            info_dic["建筑年代"] = [table.cell(2, 1).text, table.cell(2, 3).text, table.cell(2, 5).text, table.cell(2, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 3:
            info_dic["房屋栋数"] = [table.cell(3, 1).text, table.cell(3, 3).text, table.cell(3, 5).text, table.cell(3, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 4:
            info_dic["房屋结构类型"] = [table.cell(4, 1).text, table.cell(4, 3).text, table.cell(4, 5).text, table.cell(4, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 5:
            info_dic["房屋间数"] = [table.cell(5, 1).text, table.cell(5, 3).text, table.cell(5, 5).text, table.cell(5, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 6:
            info_dic["房屋尺寸"] = [table.cell(6, 1).text, table.cell(6, 3).text, table.cell(6, 5).text, table.cell(6, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 7:
            info_dic["砌块类型"] = [table.cell(7, 1).text, table.cell(7, 3).text, table.cell(7, 5).text, table.cell(7, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 8:
            info_dic["砖体粘结"] = [table.cell(8, 1).text, table.cell(8, 3).text, table.cell(8, 5).text, table.cell(8, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 9:
            info_dic["房屋墙体厚度"] = [table.cell(9, 1).text, table.cell(9, 3).text, table.cell(9, 5).text, table.cell(9, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 10:
            info_dic["房屋圈梁情况"] = [table.cell(10, 1).text, table.cell(10, 3).text, table.cell(10, 5).text, table.cell(10, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 11:
            info_dic["房屋地梁情况"] = [table.cell(11, 1).text, table.cell(11, 3).text, table.cell(11, 5).text, table.cell(11, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 12:
            info_dic["房屋构造柱情况"] = [table.cell(12, 1).text, table.cell(12, 3).text, table.cell(12, 5).text, table.cell(12, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 13:
            info_dic["场地条件"] = [table.cell(13, 1).text, table.cell(13, 3).text, table.cell(13, 5).text, table.cell(13, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 14:
            info_dic["房屋基坑深度"] = [table.cell(14, 1).text, table.cell(14, 3).text, table.cell(14, 5).text, table.cell(14, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 15:
            info_dic["基坑回填材料"] = [table.cell(15, 1).text, table.cell(15, 3).text, table.cell(15, 5).text, table.cell(15, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 16:
            info_dic["基础砌体"] = [table.cell(16, 1).text, table.cell(16, 3).text, table.cell(16, 5).text, table.cell(16, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 17:
            info_dic["基础砌筑材料"] = [table.cell(17, 1).text, table.cell(17, 3).text, table.cell(17, 5).text, table.cell(17, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 18:
            info_dic["屋盖类别"] = [table.cell(18, 1).text, table.cell(18, 3).text, table.cell(18, 5).text, table.cell(18, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 19:
            info_dic["楼板类别"] = [table.cell(19, 1).text, table.cell(19, 3).text, table.cell(19, 5).text, table.cell(19, 7).text] \
                if flag \
                else ["", "", "", ""]
        elif i == 20:
            # 有个别文档中的表格自身不规范
            try:
                info_dic["历史震害调查"] = table.cell(20, 1).text if flag else ""
            except IndexError:
                info_dic["历史震害调查"] = ""
        elif i == 21:
            try:
                info_dic["备注"] = table.cell(21, 1).text if flag else ""
            except IndexError:
                info_dic["备注"] = ""
    return info_dic


def dict_to_excel02(key_text: str, info_dic: dict, village_name: str, path_to_store: str):
    """
    将由 word02 文档中抽取的重要信息字典生成 excel02
    :param key_text: word02中有关经纬度、调查人、联系方式、日期的重要信息
    :param info_dic: 由 get_excel02_dict 函数从 word02 文档中抽取的重要信息字典
    :param village_name: 该村村名
    :param path_to_store: 目的文件存储路径
    """
    excel = xlsxwriter.Workbook(path.join(path_to_store, "整体抗震性能统计表.xlsx"), {'constant_memory': False})
    format01 = excel.add_format({'align': 'center', 'valign': 'vdistributed', })
    format02 = excel.add_format({'align': 'left', 'valign': 'vdistributed', })
    sheet01 = excel.add_worksheet(name="sheet1")
    sheet01.merge_range("A1:I1", f"附件2-{village_name}农居整体抗震性能统计表（加盖公章）", format01)
    sheet01.merge_range("A2:I2", key_text, format01)
    sheet01.write_string("A3", "自然村地址")
    sheet01.merge_range("B3:C3", info_dic["自然村地址"], format02)
    sheet01.merge_range("D3:E3", "住户总数（户）", format02)
    sheet01.write_string("F3", info_dic["住户总数（户）"], format02)
    sheet01.merge_range("G3:H3", "人口总数（口）", format02)
    sheet01.write_string("I3", info_dic["人口总数（口）"], format02)
    sheet01.write_string("A4", "房屋总数（栋）", format02)
    sheet01.write_string("B4", info_dic["房屋总数（栋）"], format02)
    sheet01.write_string("C4", "家庭平均人口（口）", format02)
    sheet01.write_string("D4", info_dic["家庭平均人口（口）"], format02)
    sheet01.merge_range("E4:G4", "全村上一年经济收入（万元）", format02)
    sheet01.merge_range("H4:I4", info_dic["全村上一年经济收入（万元）"], format02)
    key_list = [
        "建筑年代", "房屋栋数", "房屋结构类型", "房屋间数", "房屋尺寸", "砌块类型", "砖体粘结", "房屋墙体厚度", "房屋圈梁情况",
        "房屋地梁情况", "房屋构造柱情况", "场地条件", "房屋基坑深度", "基坑回填材料", "基础砌体", "基础砌筑材料", "屋盖类别", "楼板类别"
    ]
    for i in range(0, len(key_list)):
        sheet01.write_string(f"A{i + 5}", key_list[i])
        sheet01.merge_range(f"B{i + 5}:C{i + 5}", info_dic[key_list[i]][0], format02)
        sheet01.merge_range(f"D{i + 5}:E{i + 5}", info_dic[key_list[i]][1], format02)
        sheet01.merge_range(f"F{i + 5}:G{i + 5}", info_dic[key_list[i]][2], format02)
        sheet01.merge_range(f"H{i + 5}:I{i + 5}", info_dic[key_list[i]][3], format02)
    sheet01.write_string("A23", "历史震害调查", format02)
    sheet01.merge_range("B23:I23", info_dic["历史震害调查"], format02)
    sheet01.write_string("A24", "备注", format02)
    sheet01.merge_range("B24:I24", info_dic["备注"], format02)
    excel.close()


def get_filepath(root_path: str, file_list=List[str]) -> List[str]:
    """
    递归查找 root_path 下的所有文件
    :param root_path: 被查找根目录
    :param file_list: 用来存储 root_path 下所有文件路径的列表，初始值赋空列表 []
    :return: 包含 root_path 下所有文件路径的列表
    """
    if file_list is None:
        file_list = []
    for file in os.listdir(root_path):
        file_path = path.join(root_path, file)
        if path.isdir(file_path):
            get_filepath(file_path, file_list)
        else:
            file_list.append(file_path)
    return file_list


def release_dir(parent_path: str, src_dir: str):
    sub_dirs = os.listdir(
        path.join(parent_path, src_dir))
    for sub_dir in sub_dirs:
        os.rename(
            path.join(parent_path, src_dir, sub_dir),
            path.join(parent_path, sub_dir))
    os.removedirs(
        path.join(parent_path, src_dir))


def repeat_dir_remove(parent_path: str, src_dir: str):
    repeat_dirs = [sub_dir for sub_dir in os.listdir(
        path.join(parent_path, src_dir)) if sub_dir == src_dir]
    if repeat_dirs:
        for repeat_dir in repeat_dirs:
            dirs = os.listdir(
                path.join(parent_path, src_dir, repeat_dir))
            for d in dirs:
                os.rename(
                    path.join(parent_path, src_dir, repeat_dir, d),
                    path.join(parent_path, src_dir, d))
            os.removedirs(
                path.join(parent_path, src_dir, repeat_dir))


def region_dir_rename(region_dir: str):
    region_dir = region_dir.replace(' ', '')
    if region_dir.find('区') != -1:
        msg = re.findall(r'^[\u4E00-\u9FA5]{1,5}区', region_dir)
        if msg:
            region_dir = msg[0]
    else:
        if 2 <= len(region_dir) <= 5:
            region_dir = region_dir + '区'
    return region_dir


def town_dir_rename(town_dir: str):
    town_dir = town_dir.replace(' ', '')
    if town_dir.find('镇') != -1:
        msg = re.findall(r'^[\u4E00-\u9FA5]{1,5}镇', town_dir)
        if msg:
            town_dir = msg[0]
    elif town_dir.find('街道') != -1:
        msg = re.findall(r'^[\u4E00-\u9FA5]{1,5}街道', town_dir)
        if msg:
            town_dir = msg[0]
    else:
        if 2 <= len(town_dir) <= 5:
            town_dir = town_dir + '镇'
    return town_dir


def village_dir_rename(village_dir: str):
    village_dir = village_dir.replace(' ', '')
    if village_dir.find('村') != -1:
        msg = re.findall(r'^[\u4E00-\u9FA5]{1,5}村', village_dir)
        if msg:
            village_dir = msg[0]
    else:
        if 2 <= len(village_dir) <= 5:
            village_dir = village_dir + '村'
        else:
            village_dir = village_dir.replace('房屋', '').replace('抗震', '').replace('调查', '').replace('表', '').replace('性能', '')
            village_dir = village_dir.replace('统计', '').replace('上报', '').replace('材料', '')
            village_dir = re.sub(r'[0-9]{1,4}.[0-9]{1,4}.[0-9]{1,4}', '', village_dir) + '村'
    if village_dir.endswith("村村"):
        village_dir = village_dir.replace("村村", "村")
    return village_dir


def clean_region_dir(region_path: str):
    # 重命名当前区
    region_dir = region_path.split('\\')[-1]
    root_path = region_path.replace(region_dir, '')
    repeat_dir_remove(root_path, region_dir)
    region_dir = region_dir_rename(region_dir)
    os.rename(region_path, path.join(root_path, region_dir))
    region_path = path.join(root_path, region_dir)
    # 重命名区中的每一个镇或者街道
    town_dirs = [d for d in os.listdir(region_path) if path.isdir(path.join(region_path, d))]
    new_town_dirs = []
    for town_dir in town_dirs:
        # 删除该路径下重名的嵌套路径
        repeat_dir_remove(region_path, town_dir)
        # 重命名镇、街道
        new_dir = town_dir_rename(town_dir)
        os.rename(path.join(region_path, town_dir), path.join(region_path, new_dir))
        new_town_dirs.append(new_dir)
    town_dirs = new_town_dirs
    # 重命名每一个村
    for town_dir in town_dirs:
        town_path = path.join(region_path, town_dir)
        village_dirs = [d for d in os.listdir(town_path) if path.isdir(path.join(town_path, d))]
        # 个别镇中把所有的村又多汇总了一层文件夹
        if len(village_dirs) == 1 and len(os.listdir(path.join(town_path, village_dirs[0]))) >= 6:
            release_dir(town_path, village_dirs[0])
            village_dirs = [d for d in os.listdir(town_path) if path.isdir(path.join(town_path, d))]
        for village_dir in village_dirs:
            village_path = path.join(town_path, village_dir)
            # 个别村中又把所有文件多汇总了一层文件夹
            substance_dirs = [d for d in os.listdir(village_path)]
            if len(substance_dirs) == 1 and len(os.listdir(path.join(village_path, substance_dirs[0]))) >= 2:
                release_dir(village_path, substance_dirs[0])
            # repeat_dir_remove(path.join(region_path, town_dir), village_dir)
            new_dir = village_dir_rename(village_dir)
            os.rename(path.join(town_path, village_dir), path.join(town_path, new_dir))


def photo_find_number(photo_name: str) -> Union[str, None]:
    """
    提取照片名中的编号信息
    :param photo_name: 照片名
    :return: 找到编号返回编号字符串，否则返回None
    """
    number_found = re.findall(r'[A-Za-z]{2,4}-[A-Za-z]{2,4}-[A-Za-z]{2,4}-[0-9]{2,3}', photo_name)
    if number_found:
        return number_found[0]
    else:
        return None


def photo_find_name(photo_name: str) -> Union[str, None]:
    """
    提取照片名中的姓名信息
    :param photo_name: 照片名
    :return: 找到姓名返回姓名字符串，否则返回None
    """
    name_find = re.findall(r'[\u4E00-\u9FA5]{2,4}', photo_name)
    if name_find:
        return name_find[0]
    else:
        return None
    pass


class VillageWord01Handle:

    @staticmethod
    def case00() -> None:
        """
        未找到任何附件一源文件时的处理方案
        """
        pass

    @staticmethod
    def case01(path_to_store: str, docx_file: str) -> None:
        """
        所有被抽查户的附件一同意存在一个文档中时的处理方案
        :param path_to_store: 目的文件存储路径
        :param docx_file: 文件名（绝对路径）
        """
        if not path.exists(path_to_store):
            os.mkdir(path_to_store)
        if docx_file.endswith('.doc'):
            docx_file = doc_to_docx(docx_file)
        elif docx_file.endswith('.wps'):
            docx_file = wps_to_docx(docx_file)
        docx = dx.Document(docx_file)
        paragraphs: List[dx.text.paragraph] = docx.paragraphs
        tables: List[dx.table.Table] = docx.tables
        info_dic = {'name': [], 'phone': [], 'date': []}
        # 寻找附件1的信息
        for paragraph in paragraphs:
            text: str = paragraph.text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace(',', '.').replace('，', '.')
            if text.find('户主') != -1:
                name_found = re.findall(r'户主 *[:：] *[\u4e00-\u9fa5]{1,3}', text)
                if name_found:
                    name: str = name_found[0]
                    name = name.split(':')[-1]
                    info_dic['name'].append(name)
                else:
                    info_dic['name'].append("")
            if text.find('联系') != -1:
                phone_found = re.findall(r'\d{8,11}', text)  # 查找电话号码 可能需要更加精细
                if phone_found:
                    phone = phone_found[0]
                    info_dic['phone'].append(phone)
                else:
                    info_dic['phone'].append("")
            if text.find('时间') != -1:
                date_found = re.findall(r'\d{4}[.年]\d{1,2}[.月]\d{1,2}日?', text)
                # 存在问题 日期如20220825 匹配不到 若修改为 r'\d{4}[.年]?\d{1,2}[.月]?\d{1,2}日?' 会有手机号码冲突
                if date_found:
                    date = date_found[0]
                    info_dic['date'].append(date)
                else:
                    info_dic['date'].append("")
        info = [
            f"户主:{info_dic['name'][i]}     联系方式:{info_dic['phone'][i]}     日期:{info_dic['date'][i]}"
            for i in range(len(info_dic['name']))
        ]
        tables_new = [
            table
            for table in tables
            if len(table._cells) >= 100 and table.cell(0, 1).text == "编号"
        ]
        docx_element_dic = {
            tables_new[i].cell(0, 2).text.replace('\n', '').replace(' ', ''): (info[i], tables_new[i])
            for i in range(len(tables_new))
        }
        name_dic = dict()
        for name, (paragraph, table) in docx_element_dic.items():
            docx = dx.Document()
            new_paragraph = docx.add_paragraph(paragraph)
            # ————————
            # trick: 将原word中的表格复制到新表格中去（https://learnku.com/python/t/52624）
            new_table = deepcopy(table)
            new_paragraph._p.addnext(new_table._element)
            # ————————
            if name not in name_dic.keys():
                name_dic[name] = 0
            if not path.exists(path.join(path_to_store, f"{name}.docx")):
                #  可能存在编号重叠错误
                docx.save(path.join(path_to_store, f"{name}.docx"))
                name_dic[name] = 0
            else:
                docx.save(path.join(path_to_store, f"{name}-{name_dic[name] + 1:02d}.docx"))
                name_dic[name] += 1
        # 寻找附件2的信息

    @staticmethod
    def case02(path_to_store: str, word01_files_ls: List[str]) -> None:
        """
        所有被抽查户的附件一分别存在一个单独的文档中时的处理方案
        :param path_to_store: 目的文件存储路径
        :param word01_files_ls: 包含文件名（绝对路径）的列表
        """
        if not path.exists(path_to_store):
            os.mkdir(path_to_store)
        name_dic = dict()
        for word01 in word01_files_ls:
            if word01.endswith('.doc'):
                word01 = doc_to_docx(word01)
            elif word01.endswith('.wps'):
                word01 = wps_to_docx(word01)
            docx = dx.Document(word01)
            name = docx.tables[0].cell(0, 2).text
            if not path.exists(path.join(path_to_store, f"{name}.docx")):
                os.rename(word01, path.join(path_to_store, f"{name}.docx"))
                name_dic[name] = 0
            else:
                os.rename(word01, path.join(path_to_store, f"{name}-{name_dic[name] + 1:02d}.docx"))
                name_dic[name] += 1
        return None


class VillageWord02Handle:

    @staticmethod
    def case00(path_to_store: str, docx_file: str) -> None:
        """
        未找到 word02 源文件时的处理方案
        """
        docx = dx.Document(docx_file)
        paragraphs: List[dx.text.paragraph] = docx.paragraphs
        tables: List[dx.table.Table] = docx.tables
        word02_text = ''
        word02_table = None
        for paragraph in paragraphs:
            text: str = paragraph.text.replace(' ', '').replace('：', ':').replace(';', ':').replace('；', ':').replace(',', '.').replace('，', '.')
            if text.find('经纬度') != -1 and text.find('调查人') != -1:
                text = text.replace('经纬度:', '|').replace('调查人:', '|').replace('联系方式:', '|').replace('日期:', '|').replace('时间:', '|')
                text = text.replace('经纬度', '|').replace('调查人', '|').replace('联系方式', '|').replace('日期', '|').replace('时间', '|')
                text_ls = text.split('|')
                coordinate, investigator, phone, date = "", "", "", ""
                for text in text_ls:
                    if re.match(r'^北纬.*东经.*$', text):
                        coordinate = text
                    elif re.match(r'^[\u4e00-\u9fa5]{1,4}$', text):
                        investigator = text
                    elif re.match(r'^\d{8,11}$', text):
                        phone = text
                    elif re.match(r'\d{4}[.年]\d{1,2}[.月]\d{1,2}日?', text):
                        date = text
                word02_text = f"经纬度:{coordinate}    调查人:{investigator}  联系方式:{phone}    日期:{date}"
                break
        for table in tables:
            if table._cells[0].text == "自然村地址":
                word02_table = table
        if word02_text and word02_table:
            docx = dx.Document()
            # 创建横向页面
            section = docx.sections[0]
            new_width, new_height, = section.page_height, section.page_width
            section.orientation = dx.enum.section.WD_ORIENT.LANDSCAPE
            section.page_width = new_width
            section.page_height = new_height
            #
            new_paragraph = docx.add_paragraph(word02_text)
            new_table = deepcopy(word02_table)
            new_paragraph._p.addnext(new_table._element)
            docx.save(path.join(path_to_store, "附件2-整体抗震性能统计表.docx"))
            # os.remove(docx_file)

    @staticmethod
    def case01(path_to_store: str, word02_file: str) -> None:
        """
        找到 word02 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param word02_file: 文件名（绝对路径）
        """
        os.rename(word02_file, path.join(path_to_store, "附件2-整体抗震性能统计表.docx"))


class VillageExcel01Handle:

    @staticmethod
    def case00(path_to_store, region_name: str, town_name: str, village_name: str, excel01_root_path: str) -> None:
        """
        未找到 excel01 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param region_name: 该村所在区名
        :param town_name: 该村所在镇名
        :param village_name: 该村村名
        :param excel01_root_path: 存储 word01 的上级根目录
        """
        files = os.listdir(excel01_root_path)
        excel = xlsxwriter.Workbook(path.join(path_to_store, "单体抗震性能调查表.xlsx"), {'constant_memory': False})
        sheet01 = excel.add_worksheet(name="sheet1")
        format01 = excel.add_format({'align': 'left', 'valign': 'vdistributed', })
        names = [
            "编号", "北纬", "东经", "区", "乡镇", "村庄", "户主", "户主联系方式", "建筑年代", "层数", "结构类型", "建筑高度",
            "建筑宽度", "建筑长度", "常住人口", "平面", "立面", "房屋间数", "场地条件", "基坑深度", "基坑回填材料", "基础材料",
            "基础砌筑砂浆材料", "墙体砌块材料", "墙体砂浆材料", "外墙厚度", "内墙厚度", "烟道", "女儿墙", "上部圈梁", "上部圈梁闭合",
            "基础圈梁", "基础圈梁闭合", "构造柱", "屋盖类别", "楼板类别", "墙体歪闪", "墙体腐蚀", "墙体裂缝", "基础沉降", "屋盖情况",
            "填表人", "填表人联系方式"
        ]
        heads = [
            'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z',
            'AA', 'AB', 'AC', 'AD', 'AE', 'AF', 'AG', 'AH', 'AI', 'AJ', 'AK', 'AL', 'AM', 'AN', 'AO', 'AP', 'AQ'
        ]
        for name, head in zip(names, heads):
            sheet01.write_string(head + '1', name, format01)
        location = {"region_name": region_name, "town_name": town_name, "village_name": village_name}
        for i, file in enumerate(files):
            docx = dx.Document(path.join(excel01_root_path, file))
            info_text = ""
            for paragraph in docx.paragraphs:
                text: str = paragraph.text
                if text.find("户主") != -1 and text.find("联系方式") != -1:
                    info_text = text
            name_found = re.findall(r"[\u4E00-\u9FA5]{2,4}", info_text)
            phone_found = re.findall(r"\d{8,11}", info_text)
            name_found.remove("户主")
            name_found.remove("联系方式")
            if "时间" in name_found:
                name_found.remove("时间")
            elif "日期" in name_found:
                name_found.remove("日期")
            info_dic = get_excel01_dict(location, name_found[0] if name_found else "", phone_found[0] if phone_found else "", docx.tables[0])
            for name, head in zip(names, heads):
                sheet01.write_string(head + f'{i + 2}', info_dic[name], format01)
        excel.close()

    @staticmethod
    def case01(path_to_store: str, excel01_file: str) -> None:
        """
        找到 excel01 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param excel01_file: 文件名（绝对路径）
        """
        if excel01_file.endswith('.xls'):
            excel01_file = xls_to_xlsx(excel01_file)
        os.rename(excel01_file, path.join(path_to_store, "单体抗震性能调查表.xlsx"))


class VillageExcel02Handle:

    @staticmethod
    def case00(path_to_store: str, village_name: str, excel02_file: str) -> None:
        """
        未找到 excel02 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param village_name: 该村村名
        :param excel02_file: 文件名（绝对路径）
        """
        if excel02_file.endswith('doc'):
            excel02_file = doc_to_docx(excel02_file)
        elif excel02_file.endswith('wps'):
            excel02_file = wps_to_docx(excel02_file)
        docx = dx.Document(excel02_file)
        key_text = ""
        for paragraph in docx.paragraphs:
            if re.match(".*经纬度.*调查人.*联系方式.*日期.*", paragraph.text):
                key_text = paragraph.text
                break
        info_dic = get_excel02_dict(docx.tables[0])
        dict_to_excel02(key_text, info_dic, village_name, path_to_store)

    @staticmethod
    def case01(path_to_store: str, excel02_file: str) -> None:
        """
        找到 excel02 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param excel02_file: 文件名（绝对路径）
        """
        if excel02_file.endswith('.xls'):
            excel02_file = xls_to_xlsx(excel02_file)
        os.rename(excel02_file, path.join(path_to_store, "整体抗震性能统计表.xlsx"))


class VillagePhotosHandle:

    @staticmethod
    def case00() -> None:
        """
        未找到照片时的处理方案
        """
        pass

    @staticmethod
    def case01(path_to_store: str, photo_files_ls: List[str]) -> None:
        """
        找到照片时的处理方案
        :param path_to_store: 目的文件存储路径
        :param photo_files_ls: 包照片文件名（绝对路径）的列表
        """
        photo_names_ls = [file.split('\\')[-1] for file in photo_files_ls]
        photo_new_names_ls = []
        for photo_name in photo_names_ls:
            number_found = photo_find_number(photo_name)
            name_found = photo_find_name(photo_name)
            if number_found:
                photo_new_names_ls.append(number_found)
            elif name_found:
                photo_new_names_ls.append(name_found)
            else:
                photo_new_names_ls.append(photo_name)
        serial_number_dic = {}
        for i in range(len(photo_files_ls)):
            photo_name = photo_new_names_ls[i]
            photo_file = photo_files_ls[i]
            if not (
                    re.match(r'^[A-Za-z]{2,4}-[A-Za-z]{2,4}-[A-Za-z]{2,4}-[0-9]{2,3}$', photo_name)
                    or re.match(r'^[\u4E00-\u9FA5]{2,4}$', photo_name)
            ):
                if not path.exists(path.join(path_to_store, "未知")):
                    os.mkdir(path.join(path_to_store, "未知"))
                os.rename(photo_file, path.join(path_to_store, "未知", photo_name))
            else:
                if photo_name not in serial_number_dic.keys():
                    if not path.exists(path.join(path_to_store, photo_name)):
                        os.mkdir(path.join(path_to_store, photo_name))
                    serial_number_dic[photo_name] = 1
                else:
                    serial_number_dic[photo_name] += 1
                postfix = '.' + photo_file.split('.')[-1]
                os.rename(photo_file, path.join(path_to_store, photo_name, f"{photo_name}-{serial_number_dic[photo_name]:02d}{postfix}"))


class TownWord02Handle:

    @staticmethod
    def case00() -> None:
        pass

    @staticmethod
    def case01(path_to_store: str, region_name: str, town_name: str, word02_file: str) -> None:
        """
        找到 word02 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param region_name: 区名
        :param town_name: 镇名
        :param word02_file: 文件名（绝对路径）
        """
        os.rename(word02_file, path.join(path_to_store, f"{region_name}-{town_name}-整体抗震性能统计表.docx"))


class TownExcel01Handle:

    @staticmethod
    def case00() -> None:
        """
        未找到 excel01 源文件时的处理方案
        """
        pass

    @staticmethod
    def case01(path_to_store: str, region_name: str, town_name: str, excel01_file: str) -> None:
        """
        找到 excel01 源文件时的处理方案
        :param path_to_store: 目的文件存储路径
        :param region_name: 区名
        :param town_name: 镇名
        :param excel01_file: 文件名（绝对路径）
        """
        os.rename(excel01_file, path.join(path_to_store, f"{region_name}-{town_name}-单体抗震性能调查表.xlsx"))
